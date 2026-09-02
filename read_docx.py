from docx import Document

doc = Document(r"C:\Users\shaio\big-steps-travel-demo\ITINERARY.docx")

output = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        output.append(text)

for table in doc.tables:
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells if c.text.strip()]
        if cells:
            output.append(" | ".join(cells))

with open(r"C:\Users\shaio\big-steps-travel-demo\ITINERARY.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(output))

print(f"Done: {len(output)} lines extracted")
